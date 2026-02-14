import os
from docx import Document
from docx.shared import Inches

def convert_md_to_docx(md_folder, docx_folder):
    """
    Převede všechny Markdown soubory z md_folder do DOCX formátu
    a uloží je do docx_folder.
    """
    if not os.path.exists(docx_folder):
        os.makedirs(docx_folder)

    for filename in os.listdir(md_folder):
        if filename.endswith(".md"):
            md_filepath = os.path.join(md_folder, filename)
            docx_filename = filename.replace(".md", ".docx")
            docx_filepath = os.path.join(docx_folder, docx_filename)

            document = Document()

            with open(md_filepath, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line.startswith('# '):
                        document.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        document.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        document.add_heading(line[4:], level=3)
                    elif line: # Jakýkoli jiný text přidat jako odstavec
                        document.add_paragraph(line)
            
            document.save(docx_filepath)
            print(f"Převedeno: {filename} -> {docx_filename}")

if __name__ == "__main__":
    MD_FOLDER = "Generated descriptions"
    DOCX_FOLDER = "Generated descriptions - doc"

    convert_md_to_docx(MD_FOLDER, DOCX_FOLDER)
    print(f"\nKonverze dokončena. Soubory jsou uloženy ve složce: {DOCX_FOLDER}")
